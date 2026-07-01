"""
Generate Sprinkler-8 Contractor Brief Rev 2 as a .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── open the original so we inherit its styles ──────────────────
orig = Document(r'D:\Projects_One\_Claude\Sprinkler Controller\CircuitBoard\schematics\ESP32_Relay_Board_Contractor_Brief.docx')

doc = Document()

# page margins to match A4
for sec in doc.sections:
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin  = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)

# ── style helpers ────────────────────────────────────────────────
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    return p

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def note(text):
    """Indented italic note paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x40, 0x00)

def table_row(tbl, cells, bold_first=False):
    row = tbl.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        run = cell.paragraphs[0].add_run(text)
        if bold_first and i == 0:
            run.bold = True
    return row

def shade_row(row, hex_color='D0D8E8'):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

def make_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_row = tbl.rows[0]
    shade_row(hdr_row, 'D0D8E8')
    for cell, hdr in zip(hdr_row.cells, headers):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
    for row_data in rows:
        row = tbl.add_row()
        for cell, text in zip(row.cells, row_data):
            cell.paragraphs[0].add_run(text)
    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = Cm(w)
    return tbl

# ════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ESP32 8-Channel WiFi Relay Controller')
r.bold = True
r.font.size = Pt(20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Hobby Board  │  Contractor PCB Design Brief  │  Rev 2.0  │  June 2026')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

# ── intro ────────────────────────────────────────────────────────
para('This document is a complete design brief for a custom ESP32-based printed circuit board. '
     'It is a personal hobby electronics project. The board is intended for JLCPCB PCB assembly. '
     'No changes to the design should be made without consulting the board owner first.')

para('This brief provides everything a PCB designer needs to complete the schematic and layout: '
     'functional block descriptions, component specifications with LCSC part numbers, a suggested '
     'bill of materials, PCB specifications, and a prototype verification sequence.')

# ════════════════════════════════════════════════════════════════
# 1. BOARD OVERVIEW
# ════════════════════════════════════════════════════════════════
h1('1.  Board Overview')

make_table(
    ['Attribute', 'Value'],
    [
        ['Microcontroller',   'ESP32-WROOM-32E (Xtensa dual-core 240 MHz, integrated 4 MB flash, 802.11 b/g/n WiFi, Bluetooth)'],
        ['Relay Outputs',     '8 × SPST-NO electromechanical relays, rated 5 A / 250 VAC contacts'],
        ['Relay Logic',       'Active HIGH — relay energises when the corresponding ESP32 GPIO is driven HIGH'],
        ['AC Power Input',    'Single 2-position 3.5 mm screw terminal (J_PWR). L1 on pin 1, L2 on pin 2. 24 VAC.'],
        ['Board Supply Rails','+ 5 V DC (relay coils, LED drivers, CH340C) and +3.3 V DC (ESP32 and logic)'],
        ['USB Interface',     'USB Type-C, CH340C USB-to-UART bridge with auto-reset circuit (Q1, Q2)'],
        ['Output Indicators', '8 × red SMD LEDs (one per relay channel) + 1 × green status LED (GPIO23)'],
        ['Output Protection', 'MOV (metal oxide varistor) across each relay output; ULN2803A internal flyback diodes protect relay coils'],
        ['User Controls',     'DIP switch – 8 pos., raised type, labeled “Zones”; Rotary switch – 4 pos., labeled “Time” (15/30/45/60 min); '
                              'Illuminated push button, labeled “Man. Start”; BOOT button (IO0); RST button (EN)'],
        ['PCB',               '2-layer FR4, 1.6 mm, ENIG surface finish, green solder mask'],
        ['Board Size (target)','Approximately 200 mm \xd7 80 mm. Electronics zone ~175 mm; 25 mm case-divider zone at terminal end.'],
        ['Mounting',          'Four M3 mounting holes at board corners for standoff mounting inside enclosure'],
        ['Enclosure',         'Split plastic enclosure. Top section covers electronics. Bottom/front section exposes terminal strip and user controls.'],
        ['Target Assembly',   'JLCPCB PCBA — all SMD components from LCSC catalogue; THT and panel-mount parts hand-assembled'],
    ],
    col_widths=[5.0, 10.5]
)

# ════════════════════════════════════════════════════════════════
# 2. CIRCUIT BLOCK DESCRIPTIONS
# ════════════════════════════════════════════════════════════════
h1('2.  Circuit Block Descriptions')

# ── 2.1 Power Supply ────────────────────────────────────────────
h2('2.1  Power Supply')
para('The board is powered from a 24 VAC input. The power supply block converts this to regulated '
     '+5 V DC and +3.3 V DC for the on-board electronics.')

h3('Input Connector')
bullet('Single 2-position screw terminal (J_PWR), 3.5 mm pitch, rated ≥10 A / 300 V. '
       'Pin 1 labelled L1 (24 VAC Line 1); Pin 2 labelled L2 (24 VAC Line 2). '
       'This is the only AC power input — no barrel jack, no jumper.')

h3('Fusing — Three Independent Fuses')
para('Three resettable PTC polyfuses protect separate current paths; a fault in one path '
     'does not affect the others.')
bullet('F1 — 500 mA PTC on 24 VAC L1, before the bridge rectifier. '
       'Protects the electronics (power supply ICs, relay coils).')
bullet('F2 — 500 mA PTC on 24 VAC L2 return, before the bridge rectifier. '
       'Both AC legs are independently fused at 500 mA for symmetrical circuit protection.')
bullet('F3 — 1 A PTC on the 24 VAC L1 path to the relay contact commons bus. '
       'Completely independent of F1. A wiring fault on the output side blows F3 '
       'without affecting the electronics (F1 and the ESP32 remain powered).')

h3('AC to DC Conversion')
bullet('Bridge rectifier: MB10F SMD full-wave bridge (LCSC C441951). '
       '24 VAC in → approximately +33 V DC out. 1 A / 1000 V rated.')
bullet('Filter capacitor: 470 µF / 50 V electrolytic, placed directly at rectifier output.')

h3('+5 V Rail — LM2596S-5.0 Buck Converter')
bullet('IC: LM2596S-5.0, TO-263-5 package (LCSC C2621). '
       'Input up to 40 V, fixed 5 V output, 3 A rated.')
bullet('External components per datasheet: 100 µH inductor (rated 1 A+), '
       '220 µF / 10 V output electrolytic capacitor, 1N5822 Schottky catch diode.')

h3('+3.3 V Rail — AMS1117-3.3 LDO')
bullet('IC: AMS1117-3.3, SOT-223 package (LCSC C6186). Powered from the +5 V rail.')
bullet('Input capacitor: 10 µF ceramic. Output capacitor: 22 µF ceramic.')

# ── 2.2 Microcontroller ─────────────────────────────────────────
h2('2.2  Microcontroller — ESP32-WROOM-32E')
para('The ESP32-WROOM-32E is a fully integrated module containing the ESP32 SoC, 4 MB flash, '
     '40 MHz crystal, and a PCB trace antenna. Use LCSC C701341.')
bullet('Supply: +3.3 V. Decouple with 10 µF + 100 nF ceramic capacitors at the module power pins.')
bullet('EN pin (chip enable): Pull up to 3.3 V via 10 kΩ. RST push button pulls EN to GND.')
bullet('IO0 pin (boot mode): Pull up to 3.3 V via 10 kΩ. BOOT push button pulls IO0 to GND. '
       'IO0 = LOW at reset enters download mode for firmware flashing.')
bullet('GPIO12 strapping pin — CRITICAL: GPIO12 is NOT used as a relay output. '
       'Connect a 10 kΩ resistor from GPIO12 to GND and leave the pin otherwise unconnected. '
       'This pin controls SDIO voltage selection at boot — if it reads HIGH the module selects '
       '1.8 V flash voltage and will not start. The pull-down resistor is mandatory. '
       'Do not connect GPIO12 to the ULN2803A or any other driven signal.')
bullet('Antenna keep-out: Maintain a copper-free zone of at least 15 mm around the on-module PCB '
       'trace antenna. No copper on any layer, no vias, no components in this zone.')

h3('GPIO Assignments — Do Not Deviate')
make_table(
    ['GPIO', 'Function', 'Direction', 'Notes'],
    [
        ['GPIO4',  'Relay Channel 1', 'Output', '→ ULN2803A input 1'],
        ['GPIO13', 'Relay Channel 2', 'Output', '→ ULN2803A input 2'],
        ['GPIO14', 'Relay Channel 3', 'Output', '→ ULN2803A input 3'],
        ['GPIO25', 'Relay Channel 4', 'Output', '→ ULN2803A input 4'],
        ['GPIO26', 'Relay Channel 5', 'Output', '→ ULN2803A input 5'],
        ['GPIO27', 'Relay Channel 6', 'Output', '→ ULN2803A input 6'],
        ['GPIO32', 'Relay Channel 7', 'Output', '→ ULN2803A input 7'],
        ['GPIO33', 'Relay Channel 8', 'Output', '→ ULN2803A input 8'],
        ['GPIO12', 'Strapping pin — NOT a relay output', 'N/A', '10 kΩ pull-down to GND only. Do not connect to ULN2803A or any driven net. See Section 2.2 note.'],
        ['GPIO23', 'Status LED',      'Output', 'Active HIGH → 330 Ω → green LED anode → cathode to GND'],
        ['GPIO1 / TXD0', 'UART TX',   'Output', '→ CH340C RXD pin'],
        ['GPIO3 / RXD0', 'UART RX',   'Input',  '← CH340C TXD pin'],
        ['GPIO0 / IO0',  'Boot select','I/O',    '10 kΩ pull-up to 3.3 V. BOOT button to GND. Auto-reset uses this pin.'],
        ['EN',           'Chip enable/reset', 'Input', '10 kΩ pull-up to 3.3 V. RST button to GND. Auto-reset uses this pin.'],
        ['GPIO16', 'Rotary switch BCD bit 0 (Time)', 'Input', '← SW_ROT BCD output A. 10 kΩ pull-up to 3.3 V. Active LOW.'],
        ['GPIO17', 'Rotary switch BCD bit 1 (Time)', 'Input', '← SW_ROT BCD output B. 10 kΩ pull-up to 3.3 V. Active LOW.'],
        ['GPIO18', 'Manual Start button input', 'Input', '← SW_START contact. Active LOW, 10 kΩ pull-up to 3.3 V.'],
        ['GPIO19', 'Manual Start button LED',   'Output','→ 330 Ω → SW_START LED anode. Active HIGH.'],
        ['GPIO21', 'I²C SDA — MCP23008',        'I/O',   'SDA to MCP23008 pin 2. 4.7 kΩ pull-up to 3.3 V.'],
        ['GPIO22', 'I²C SCL — MCP23008',        'Output','SCL to MCP23008 pin 1. 4.7 kΩ pull-up to 3.3 V.'],
        ['(MCP23008)', 'DIP switch inputs Zone 1–8', 'Input', 'All 8 DIP switch inputs routed to MCP23008 GP0–GP7 (not direct ESP32 GPIOs). See Section 2.6.'],
    ],
    col_widths=[2.8, 4.0, 2.0, 6.7]
)
note('All GPIO assignments are final. DIP switch inputs are routed through the MCP23008 I²C GPIO expander — '
     'do NOT connect DIP switch lines directly to the ESP32. SDA and SCL require 4.7 kΩ pull-ups to 3.3 V.')

# ── 2.3 USB-C ───────────────────────────────────────────────────
h2('2.3  USB-C Programming Interface — CH340C')
para('The CH340C provides USB-to-UART conversion for firmware programming and serial monitoring. '
     'It has an internal oscillator and requires no external crystal.')
bullet('IC: CH340C, SOP-16 package (LCSC C84681).')
bullet('USB connector: USB Type-C receptacle, 16-pin SMD (LCSC C165948).')
bullet('CC1 and CC2 pins: Each connected to GND through a 5.1 kΩ resistor. Mandatory for USB-C host compatibility.')
bullet('ESD protection: USBLC6-2SC6 SOT-23-6 TVS diode array (LCSC C2827693) between the USB-C D+/D− pins and the CH340C.')
bullet('VBUS polyfuse: 500 mA resettable PTC fuse in series with the USB VBUS line (LCSC C386618).')
bullet('Decoupling: 100 nF ceramic capacitor at CH340C VCC pin.')

h3('Auto-Reset Circuit')
para('Two MMBT3904 NPN transistors (LCSC C20526) allow the programming tool to automatically '
     'place the ESP32 into bootloader mode without manually pressing BOOT and RST.')
bullet('Q1 (RTS → EN): Base connected to CH340C RTS# pin through 10 kΩ resistor. '
       'Collector connected to ESP32 EN pin. Add 100 nF base-to-GND filter capacitor.')
bullet('Q2 (DTR → IO0): Base connected to CH340C DTR# pin through 10 kΩ resistor. '
       'Collector connected to ESP32 IO0 pin. Add 100 nF base-to-GND filter capacitor.')

# ── 2.4 Relay Driver and Output Block ───────────────────────────
h2('2.4  Relay Driver and Output Block')

h3('ULN2803A — Relay Driver IC')
para('The ULN2803A (LCSC C56026, SOP-18) is an 8-channel Darlington transistor array. '
     'It drives the relay coils from ESP32 GPIO signals (3.3 V logic compatible).')
bullet('Inputs 1–8 (pins 1–8): Connected directly to ESP32 GPIO4, 13, 14, 25, 26, 27, 32, 33 respectively (relay channels 1–8). '
       'GPIO12 is NOT connected here — it is a dedicated strapping resistor only (see Section 2.2).')
bullet('COM pin (pin 10): MUST connect to +5 V. This is the return rail for the internal flyback diodes. '
       'Leaving it floating will allow voltage spikes to damage the device.')
bullet('Collector outputs (pins 11–18): Each connected to the relay coil negative terminal.')
bullet('GND (pin 9): Connect to GND.')

h3('Relays — HF46F-G/5-HS1')
bullet('Type: SPST-NO (normally open). Contact rating: 5 A / 250 VAC.')
bullet('Coil: 5 V DC, G5NB-compatible SMD footprint.')
bullet('LCSC part number: C165255. Do not substitute — C128511 is a capacitor, not a relay.')
bullet('Coil positive terminal: +5 V. Coil negative terminal: ULN2803A collector output.')
bullet('Relay NO contact: connects to the corresponding zone output screw terminal (J_ZONES, Z1–Z8).')
bullet('Relay contact commons bus: All 8 relay COM contacts are bussed together and connected '
       'to 24 VAC L1 through fuse F3. This bus is entirely independent of the electronic power supply.')

h3('Zone Output and Common Terminals')
bullet('J_ZONES: Single 8-position screw terminal block, 3.5 mm pitch, rated ≥10 A / 300 V. '
       'Positions labeled Z1–Z8. Each position connects to the NO contact of its relay.')
bullet('J_COM: One 2-position screw terminal, 3.5 mm pitch, labeled “COM / L2”. '
       'Both pins connect to the same node — 24 VAC L2 — which also connects to J_PWR pin 2. '
       'Two screw positions are provided for convenient daisy-chain wiring from the load.')
bullet('J_COM is electrically identical to the L2 terminal of J_PWR. '
       'No additional conversion or relay contact is in this path.')
bullet('Place J_ZONES and J_COM together along the terminal-zone edge of the board for clean field wiring.')

# ── 2.5 User Controls (NEW) ─────────────────────────────────────
h2('2.5  User Controls')

h3('DIP Switch — Zone Select (SW_DIP)')
bullet('8-position DIP switch, raised (piano-key) style. '
       'Switches must protrude enough to operate by hand without a tool.')
bullet('Silkscreen: Print “Zones” above the switch body; '
       'print position numbers 1–8 below each switch.')
bullet('Switch outputs connect to MCP23008 GP0–GP7 (not directly to ESP32 GPIOs). '
       'The MCP23008 internal pull-ups are enabled in firmware — no external pull-up resistors '
       'required on the switch lines. Switch closes to GND when activated. Active LOW.')
bullet('See Section 2.6 for MCP23008 wiring and addressing.')

h3('Rotary Switch — Time Select (SW_ROT)')
bullet('4-position binary-coded (BCD) rotary switch with clear mechanical detents.')
bullet('Silkscreen: Print “Time” above the switch. '
       'Print the time value at each detent position: 15, 30, 45, 60 (or add “min” if space allows).')
bullet('Position 0 = 15 minutes; Position 1 = 30 minutes; Position 2 = 45 minutes; Position 3 = 60 minutes.')
bullet('2-bit BCD output — two GPIO lines (TBD) plus GND. '
       'Confirm active HIGH or active LOW with the selected rotary part datasheet.')
bullet('Panel-mount style with knob preferred over PCB-mount with slot for ease of use.')

h3('Push Button — Manual Start (SW_START)')
bullet('22 mm panel-mount illuminated momentary push button, green LED.')
bullet('Silkscreen: Print “Man. Start” adjacent to the button.')
bullet('Button contact: Normally open, momentary. Active LOW (10 kΩ pull-up to 3.3 V). '
       'Connected to a TBD ESP32 GPIO.')
bullet('Button LED: Driven from a TBD ESP32 GPIO through a 330 Ω series resistor. '
       'Illuminated when a manual run is active.')
bullet('Confirm mounting hole size and thread type with the selected button before routing.')

# ── 2.6 I2C GPIO Expander ───────────────────────────────────────
h2('2.6  I²C GPIO Expander — MCP23008')
para('The MCP23008 is an 8-channel I²C GPIO expander used to read all 8 DIP switch inputs '
     'over a 2-wire I²C bus, freeing ESP32 GPIOs for other functions.')

h3('Connections')
bullet('VDD: 3.3 V. GND: GND.')
bullet('SDA (pin 2): Connected to ESP32 GPIO21. 4.7 kΩ pull-up resistor to 3.3 V on this line.')
bullet('SCL (pin 1): Connected to ESP32 GPIO22. 4.7 kΩ pull-up resistor to 3.3 V on this line.')
bullet('RESET (pin 6): Connect to 3.3 V. Active LOW — tie HIGH for normal operation.')
bullet('Address pins A0, A1, A2 (pins 5, 4, 3): All connected to GND. '
       'This sets the I²C address to 0x20.')
bullet('GP0–GP7 (pins 10–17): Connected to DIP switch positions 1–8 respectively. '
       'Other side of each switch connects to GND. '
       'MCP23008 internal pull-ups are enabled via the GPPU register (0x06 = 0xFF) in firmware — '
       'no external pull-up resistors needed on GP0–GP7.')

h3('Firmware Operation')
bullet('Configure all 8 pins as inputs: write 0xFF to IODIR register (0x00).')
bullet('Enable internal pull-ups: write 0xFF to GPPU register (0x06).')
bullet('Read all 8 switch states in one I²C transaction: read GPIO register (0x09). '
       'Returns a single byte — bit 0 = Zone 1, bit 1 = Zone 2, …, bit 7 = Zone 8. '
       'A bit reading 0 means that switch is ON (closed to GND).')
bullet('Suggested read interval: poll every 200 ms. DIP switch state is not time-critical.')

h3('Package')
bullet('Use SOIC-18 (surface-mount) for JLCPCB PCBA compatibility. '
       'Search LCSC for "MCP23008" — verify the SOIC-18 package before ordering. '
       'DIP-18 is through-hole only and requires hand assembly.')

# ── 2.7 Indicator LEDs ──────────────────────────────────────────
h2('2.7  Indicator LEDs')
bullet('Per-channel LEDs \xd7 8 (red, 0805, LCSC C84256): One per relay channel. '
       'Connected between +5 V and the ULN2803A collector output net through a 470 Ω series resistor. '
       'LED illuminates when the relay is energised.')
bullet('Status LED \xd7 1 (green, 0805, LCSC C72043): Driven by GPIO23 through a 330 Ω '
       'series resistor to GND. Active HIGH.')

# ── 2.8 Output Protection ───────────────────────────────────────
h2('2.8  Output Protection — MOV Snubbers')
para('Inductive loads generate voltage spikes when switched off. '
     'A metal oxide varistor (MOV) across each relay output clamps these spikes.')
bullet('One MOV per channel, connected across the relay NO output terminal and the COM terminal '
       '— i.e., directly across the load terminals in the field.')
bullet('Rating: 47 V clamping voltage. Use TVR05470 or equivalent 5 mm disc MOV. Search LCSC for ‘MOV 47V disc’.')
bullet('Note: The ULN2803A internal flyback diodes separately protect the +5 V rail from relay coil back-EMF. '
       'Both protections are required.')

# ════════════════════════════════════════════════════════════════
# 3. SUGGESTED BILL OF MATERIALS
# ════════════════════════════════════════════════════════════════
h1('3.  Suggested Bill of Materials')
para('All components are sourced from LCSC where part numbers are given. '
     'Verify all part numbers and footprints in EasyEDA before ordering. '
     'Components marked “Verify LCSC” require a datasheet search to confirm the correct part. '
     'Where LCSC Basic parts exist, prefer them to minimise JLCPCB extended component surcharges.')
doc.add_paragraph()

make_table(
    ['Component / Description', 'LCSC Part', 'Package', 'Qty', 'Assembly', 'Notes'],
    [
        # ICs
        ['ESP32-WROOM-32E module',      'C701341',  'Module',    '1', 'SMD/Reflow', '4 MB flash min.'],
        ['CH340C USB-UART',             'C84681',   'SOP-16',    '1', 'SMD/Reflow', 'Internal osc. — no crystal'],
        ['AMS1117-3.3 LDO',            'C6186',    'SOT-223',   '1', 'SMD/Reflow', '3.3 V, 1 A'],
        ['LM2596S-5.0 Buck Converter', 'C2621',    'TO-263-5',  '1', 'SMD/Reflow', '5 V / 3 A, 40 V input max'],
        ['MB10F Bridge Rectifier',     'C441951',  'MBS',       '1', 'SMD/Reflow', '1 A / 1000 V'],
        ['ULN2803A Darlington Array',  'C56026',   'SOP-18',    '1', 'SMD/Reflow', '8-ch, internal flyback diodes. COM pin → +5 V'],
        ['HF46F-G/5-HS1 Relay 5 V SMD','C165255', 'SMD',       '8', 'SMD/Reflow', 'SPST-NO, G5NB footprint. Use ONLY C165255.'],
        ['USBLC6-2SC6 USB ESD',        'C2827693', 'SOT-23-6',  '1', 'SMD/Reflow', 'D+ / D− TVS protection'],
        ['MMBT3904 NPN Transistor',    'C20526',   'SOT-23',    '2', 'SMD/Reflow', 'Auto-reset circuit (Q1 and Q2)'],
        ['MCP23008 I²C GPIO Expander',  'Verify LCSC — search "MCP23008"', 'SOIC-18', '1', 'SMD/Reflow', 'Reads all 8 DIP switch inputs over I²C. Address 0x20 (A0/A1/A2 to GND). Use SOIC-18 for SMD assembly.'],
        # Connectors
        ['USB Type-C Receptacle 16-pin',    'C165948',      'SMD 16-pin',    '1',  'SMD/Reflow', 'Verify footprint before ordering'],
        ['2-position screw terminal 3.5 mm','C8465',        'THT 3.5 mm',    '3',  'Hand',       'J_PWR (power in) \xd7 1; J_COM (common return) \xd7 2'],
        ['8-position screw terminal 3.5 mm','C8465 \xd7 4 ganged or 8-pos block', 'THT 3.5 mm', '1', 'Hand', 'J_ZONES — zone outputs Z1–Z8. Use ganged 2-pin blocks or a single 8-position block.'],
        # New controls
        ['8-position DIP switch, raised/piano-key', 'Verify LCSC — search “8P DIP raised”', 'THT/SMD', '1', 'Hand', 'SW_DIP. Switches must protrude for hand operation. Confirm height clears case lid.'],
        ['4-position BCD rotary switch',  'Verify LCSC — search “BCD rotary 4 position”', 'Panel-mount', '1', 'Hand', 'SW_ROT. 4 clear detents. Panel-mount with knob preferred.'],
        ['22 mm illuminated push button, green', 'Verify LCSC — search “22mm push button green LED”', 'Panel-mount', '1', 'Hand', 'SW_START. Momentary NO, green LED, 3.3 V or 5 V LED. Confirm thread/mounting hole.'],
        # Protection
        ['MOV 47 V disc',              'Verify LCSC',  '5 mm disc',     '8', 'Hand',       'Across each relay output. Search “MOV 47V”.'],
        # Fuses
        ['500 mA PTC Polyfuse',         'Verify LCSC',  'THT or SMD',    '2', 'Reflow/Hand','F1 (L1 circuit power) and F2 (L2 circuit power). Both 500 mA.'],
        ['1 A PTC Polyfuse',           'Verify LCSC',  'THT or SMD',    '1', 'Reflow/Hand','F3 (relay contact commons bus — zone output side). 1 A rated.'],
        ['500 mA PTC Polyfuse 0805',   'C386618',      '0805',          '1', 'SMD/Reflow', 'USB VBUS protection'],
        # Discrete semis
        ['1N5822 Schottky Diode',      'Verify LCSC',  'DO-201 / SMD',  '1', 'SMD/Reflow', 'LM2596 catch diode'],
        # LEDs
        ['Red LED 0805',               'C84256',       '0805',          '8', 'SMD/Reflow', 'Per-channel relay indicators'],
        ['Green LED 0805',             'C72043',       '0805',          '1', 'SMD/Reflow', 'System status (GPIO23)'],
        # Passives
        ['470 Ω Resistor 0805',  'Basic',        '0805',          '8', 'SMD/Reflow', 'Zone LED current limit (R8–R15)'],
        ['330 Ω Resistor 0805',  'Basic',        '0805',          '2', 'SMD/Reflow', 'Status LED (R16) + Man. Start button LED (R17)'],
        ['10 kΩ Resistor 0402',  'Basic',        '0402',          '8', 'SMD/Reflow', 'Pull-ups: EN, IO0, rotary BCD A/B, start button input (×5); GPIO12 pull-down (×1); Q1/Q2 base resistors (×2)'],
        ['4.7 kΩ Resistor 0402', 'Basic',        '0402',          '2', 'SMD/Reflow', 'I²C bus pull-ups: SDA (GPIO21) and SCL (GPIO22) to 3.3 V'],
        ['5.1 kΩ Resistor 0402', 'Basic',        '0402',          '2', 'SMD/Reflow', 'USB-C CC1 and CC2 to GND (R1, R2)'],
        ['100 nF Ceramic Cap 0402',   'Basic',        '0402',          '8', 'SMD/Reflow', 'Decoupling: all IC power pins + Q1/Q2 bases (C7–C10, C9, C10)'],
        ['10 µF Ceramic Cap 0603','Basic',        '0603',          '3', 'SMD/Reflow', 'AMS1117 in/out + ESP32 supply (C3–C5)'],
        ['22 µF Ceramic Cap 0603','Basic',        '0603',          '1', 'SMD/Reflow', 'AMS1117 output cap (C4)'],
        ['470 µF Electrolytic 50 V','Verify LCSC','Radial',        '1', 'Hand',       'Bulk filter — bridge rectifier output (C1)'],
        ['220 µF Electrolytic 10 V','Verify LCSC','Radial',        '1', 'Hand',       'LM2596 output filter (C2)'],
        ['100 µH Inductor (1 A+)','Verify LCSC', 'SMD / THT',     '1', 'Reflow/Hand','LM2596 switching inductor (L1)'],
        # Buttons
        ['BOOT Push Button 6 mm',      'C318884',      'THT 6\xd76 mm', '1', 'Hand',       'IO0 to GND (SW1)'],
        ['RST Push Button 6 mm',       'C318884',      'THT 6\xd76 mm', '1', 'Hand',       'EN to GND (SW2)'],
    ],
    col_widths=[5.5, 3.5, 2.0, 0.8, 1.8, 4.4]
)

p = doc.add_paragraph()
r = p.add_run('⚠  NOTE: Relay part number C165255 (HF46F-G/5-HS1) is the correct part. '
              'C128511 is a capacitor. C22363895 is a through-hole relay. '
              'EasyEDA search results may show these by mistake — verify the part description before placing.')
r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x20, 0x00)

# ════════════════════════════════════════════════════════════════
# 4. PCB SPECIFICATIONS
# ════════════════════════════════════════════════════════════════
h1('4.  PCB Specifications')

h2('4.1  Board Parameters')
make_table(
    ['Parameter', 'Specification'],
    [
        ['Layers',           '2 (top copper and bottom copper)'],
        ['Material',         'FR4'],
        ['Thickness',        '1.6 mm'],
        ['Surface Finish',   'ENIG (Electroless Nickel Immersion Gold) — required for USB-C SMD footprint reliability'],
        ['Copper Weight',    '1 oz both layers'],
        ['Solder Mask',      'Green, both sides'],
        ['Silkscreen',       'White, both sides. All reference designators, connector labels (Z1–Z8, L1, L2, COM/L2, Zones, Time, Man. Start), '
                             'and zone boundary line must be clearly legible.'],
        ['Target Board Size','~200 mm \xd7 80 mm — electronics zone ~175 mm, case-divider zone 25 mm at terminal end. Adjust as needed to fit components.'],
        ['Minimum Trace Width','0.15 mm general; 1.5 mm minimum for relay contact and AC input traces'],
        ['Minimum Via',      '0.3 mm drill, 0.6 mm pad'],
        ['Mounting Holes',   '4 \xd7 M3 (3.2 mm drill, 6.5 mm copper ring) at corners for standoff mounting'],
        ['Fiducial Marks',   'Minimum 3 on the top copper layer for PCBA pick-and-place alignment'],
    ],
    col_widths=[5.0, 10.5]
)

h2('4.2  Board Zones and Case Divider')
para('The board is divided into two functional zones separated by a clearly marked boundary:')
bullet('Electronics zone (approximately 175 mm of board length): '
       'Contains relays K1–K8, ESP32 module, power supply (bridge, LM2596, AMS1117), '
       'CH340C UART, ULN2803A, USB-C port, BOOT/RST buttons, and all passives. '
       'Not user-accessible in normal operation.')
bullet('Terminal and control zone (the remaining ~25 mm at the terminal end of the board): '
       'Contains J_PWR, J_ZONES, J_COM, SW_DIP, SW_ROT, and SW_START. '
       'Fully user-accessible through the bottom section of the enclosure.')
bullet('Case divider zone: Mark the boundary between the two zones with a dashed line on the '
       'silkscreen layer. No components within 5 mm of this boundary on either side. '
       'The plastic case wall sits in this gap. Route inter-zone wiring (relay contact '
       'wires to terminals) through slots or notches in the case wall.')

h2('4.3  Trace Width Requirements')
make_table(
    ['Net / Signal', 'Min Width', 'Notes'],
    [
        ['24 VAC input traces',         '1.5 mm', 'Before and after bridge rectifier; includes F1 and F2 paths'],
        ['Relay contact traces',        '1.5 mm', 'F3 fuse, commons bus, and all zone output traces — 5 A rated'],
        ['+5 V power bus',              '0.8 mm', 'Supply to relay coils, CH340C, ULN2803A'],
        ['+3.3 V power bus',            '0.5 mm', 'ESP32 and logic supply'],
        ['GND pour',                    'Copper fill', 'Flood fill on both layers; via-stitch together'],
        ['GPIO signal traces',          '0.25 mm', 'Logic signals from ESP32 to ULN2803A and user controls'],
        ['USB data lines (D+, D−)','0.2 mm', 'Matched length; keep away from switching supply'],
        ['UART (TX, RX)',               '0.2 mm', 'Keep away from the LM2596 switching loop'],
        ['I²C (SDA, SCL)',         '0.2 mm', 'If GPIO expander is used; 4.7 kΩ pull-ups to 3.3 V'],
    ],
    col_widths=[5.0, 2.5, 8.0]
)

h2('4.4  Clearance Requirements')
bullet('AC safety clearance: 3 mm minimum between any 24 VAC trace and all low-voltage DC traces.')
bullet('Relay contact traces (switched AC, NO terminal): 3 mm minimum clearance from all DC logic traces.')
bullet('General PCB clearance (non-high-voltage): 0.2 mm minimum between all other traces.')
bullet('Mounting hole clearance: 5 mm from all copper, no exceptions.')
bullet('Board edge clearance: 0.3 mm minimum for copper; 0.5 mm for components.')

h2('4.5  Ground Plane')
bullet('Flood fill GND copper on both top and bottom layers.')
bullet('Via-stitch the top and bottom GND pours together with multiple vias around the perimeter and under large copper areas.')
bullet('Do NOT flood fill GND copper under the ESP32-WROOM-32E antenna area. The keep-out zone must be copper-free on all layers.')
bullet('Use a single star-ground connection point between the AC/rectifier section and the DC logic ground to prevent ground loops.')

h2('4.6  Switching Power Supply Layout (LM2596)')
para('The LM2596 switching loop must be kept as tight and compact as possible to minimise radiated EMI.')
bullet('Switching loop: LM2596 OUTPUT pin → catch diode (D2) → inductor (L1) → output capacitor (C2) → GND. '
       'Keep this entire loop area as small as possible.')
bullet('Place the catch diode directly adjacent to the LM2596 OUTPUT pin.')
bullet('Do not route USB, UART, or GPIO traces near the LM2596 switching loop.')

h2('4.7  Component Placement Guidelines')
bullet('ESP32 module: Top side. Orient antenna toward a board edge with no copper within 15 mm of the antenna.')
bullet('Relay block: Group all 8 relays together in a row near the top edge of the electronics zone.')
bullet('J_ZONES (8-position terminal): Along the terminal-zone edge. Relay channels Z1–Z8 in order, left to right.')
bullet('J_COM: Adjacent to J_ZONES, at the same edge.')
bullet('J_PWR (2-position terminal): Adjacent to J_COM at the same edge.')
bullet('SW_DIP (“Zones”): In the terminal zone, accessible from outside the case. '
       'Orient so switch numbers 1–8 are clearly readable.')
bullet('SW_ROT (“Time”): In the terminal zone, panel-mount hole if panel-mount type. '
       'Print time values (15, 30, 45, 60) on silkscreen at each detent.')
bullet('SW_START (“Man. Start”): In the terminal zone, panel-mount hole for button body. '
       'Button face must be accessible from outside the case.')
bullet('Power supply block (bridge rectifier, LM2596, AMS1117): Keep away from the relay/terminal edge. '
       'Maintain AC safety clearance from all DC logic traces.')
bullet('USB-C connector: On the short edge of the board, accessible from outside the enclosure.')
bullet('BOOT and RST buttons: Near the ESP32 module. These are internal (electronics zone) — '
       'they do not need to be accessible through the case in normal use.')
bullet('Zone LEDs: Adjacent to their corresponding relay.')
bullet('Status LED: Near the ESP32 module, near the board edge for visibility through the case.')
bullet('MOV snubbers: As close as possible to the J_ZONES terminal connections.')

# ════════════════════════════════════════════════════════════════
# 5. ENCLOSURE
# ════════════════════════════════════════════════════════════════
h1('5.  Enclosure / Case')
para('A split plastic enclosure is required. The board owner will source or fabricate this separately; '
     'the PCB designer should ensure the board layout is compatible with the requirements below.')

bullet('Material: ABS or polycarbonate, flame-retardant grade (UL94 V-0 preferred).')
bullet('Top section: Encloses the electronics zone. Closed — no user-accessible openings except '
       'possibly a cut-out for the USB-C port (for firmware updates). '
       'Case top screws or snaps onto the PCB standoffs.')
bullet('Bottom / front section: Exposes the terminal and control zone. '
       'Screw terminals must be accessible for wire connection without removing the case. '
       'DIP switch, rotary switch, and push button must be operable without tools.')
bullet('Internal divider wall: A wall integral to the case sits in the 25 mm case-divider zone on the PCB. '
       'This wall physically separates the electronics from the user-accessible terminal/control area. '
       'The wall has slots or notches for the relay contact wiring that crosses between zones.')
bullet('Mounting: Case mounts to the M3 standoffs at the four board corners.')
bullet('Colour: Light grey or off-white preferred. No specific requirement.')

# ════════════════════════════════════════════════════════════════
# 6. DELIVERABLES REQUIRED
# ════════════════════════════════════════════════════════════════
h1('6.  Deliverables Required')
para('The following files are required from the designer for JLCPCB PCBA order submission:')
bullet('Gerber files (RS-274X format): All copper layers, solder mask (both sides), '
       'silkscreen (both sides), board outline.')
bullet('Drill file: Excellon format with separate PTH (plated) and NPTH (non-plated) files.')
bullet('BOM CSV: Component reference designator, LCSC part number, quantity, footprint, and value. '
       'Must match the JLCPCB PCBA upload template.')
bullet('Pick and Place CSV (CPL): Component reference designator, centroid X, centroid Y, '
       'rotation, and layer. Must match the JLCPCB PCBA upload template.')
bullet('EasyEDA project source files (schematic + PCB layout) as backup.')

# ════════════════════════════════════════════════════════════════
# 7. PROTOTYPE VERIFICATION SEQUENCE
# ════════════════════════════════════════════════════════════════
h1('7.  Prototype Verification Sequence')
para('The following test sequence should be followed when the first prototype boards are received:')
bullet('Step 1: Power on via USB only (no 24 VAC connected). Confirm +5 V and +3.3 V rails are present. '
       'Confirm USB enumeration and serial port recognition on the host PC.')
bullet('Step 2: Flash the relay test firmware via Arduino IDE. Confirm each relay clicks in sequence '
       'and the corresponding zone LED illuminates.')
bullet('Step 3: Connect 24 VAC. Confirm bridge rectifier output is approximately +33 V DC. '
       'Confirm relay contact output voltage on J_ZONES terminals. '
       'Confirm F1, F2, F3 fuses are all intact.')
bullet('Step 4: Test F3 independence: with 24 VAC connected, short a zone output momentarily. '
       'F3 should open; F1 should remain closed and the ESP32 should remain powered.')
bullet('Step 5: Test DIP switch (SW_DIP): confirm firmware reads each switch position correctly via serial output.')
bullet('Step 6: Test rotary switch (SW_ROT): confirm firmware reads each of the 4 time positions correctly.')
bullet('Step 7: Test Manual Start button (SW_START): confirm button triggers a run, LED illuminates, '
       'and a second press stops the run.')
bullet('Step 8: Flash the main application firmware. Confirm WiFi connection and normal operation.')

# ── end line ─────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('End of Contractor Design Brief  │  Hobby Board  │  Rev 2.0')
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.font.size = Pt(10)

# ── save ─────────────────────────────────────────────────────────
out_path = r'D:\Projects_One\_Claude\Sprinkler Controller\CircuitBoard\schematics\Hobby_Board_Contractor_Brief_Rev2.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
